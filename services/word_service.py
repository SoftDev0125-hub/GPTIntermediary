"""
Microsoft Word Service
Handles Word document operations including creation, editing, formatting, and manipulation
"""

import os
import logging
from typing import Optional, List, Dict, Any
from pathlib import Path
import platform

logger = logging.getLogger(__name__)

# Try to import python-docx, but handle gracefully if not installed
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. Word features will be limited. Install with: pip install python-docx")


class WordService:
    """Service for Microsoft Word document operations"""
    
    def __init__(self):
        self.os_type = platform.system()
        self.default_documents_dir = self._get_default_documents_dir()
        logger.info(f"Word Service initialized for {self.os_type}")
        
        if not HAS_DOCX:
            logger.warning("python-docx library not available. Some Word features may not work.")
    
    def _get_default_documents_dir(self) -> str:
        """Get the default documents directory for the OS"""
        # Use the current working directory instead of C: drive
        # This avoids creating files in C: drive
        current_dir = os.getcwd()
        # If current directory is on C: drive, use D: drive instead
        if current_dir.startswith('C:') or current_dir.startswith('c:'):
            # Use D: drive Documents folder
            if os.path.exists('D:\\Documents'):
                return 'D:\\Documents'
            # Or use D: drive root
            return 'D:\\'
        # Otherwise, use current directory
        return current_dir
    
    async def initialize(self):
        """Initialize the Word service"""
        try:
            # Ensure documents directory exists
            os.makedirs(self.default_documents_dir, exist_ok=True)
            logger.info(f"Word service initialized. Documents directory: {self.default_documents_dir}")
        except Exception as e:
            logger.error(f"Error initializing Word service: {e}")
    
    async def cleanup(self):
        """Cleanup resources"""
        logger.info("Word service cleanup completed")
    
    async def create_document(
        self,
        file_path: str,
        content: Optional[str] = None,
        title: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create a new Word document
        
        Args:
            file_path: Path where the document should be saved
            content: Optional initial content for the document
            title: Optional document title
        
        Returns:
            Dictionary with success status and file path
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed. Install with: pip install python-docx"
            }
        
        try:
            # Normalize the file path
            file_path = os.path.normpath(file_path)
            
            # Ensure directory exists
            file_dir = os.path.dirname(file_path) or self.default_documents_dir
            if file_dir:
                os.makedirs(file_dir, exist_ok=True)
            
            # Create new document
            doc = Document()
            
            # Add title if provided
            if title:
                title_para = doc.add_heading(title, level=1)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add content if provided
            if content:
                doc.add_paragraph(content)
            
            # Save document
            doc.save(file_path)
            
            logger.info(f"Created Word document: {file_path}")
            return {
                "success": True,
                "file_path": file_path,
                "message": f"Document created successfully at {file_path}"
            }
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def open_document(self, file_path: str) -> Dict[str, Any]:
        """
        Open an existing Word document
        
        Args:
            file_path: Path to the document
        
        Returns:
            Dictionary with document information
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            
            # Extract document information
            paragraphs = [para.text for para in doc.paragraphs]
            
            return {
                "success": True,
                "file_path": file_path,
                "paragraph_count": len(paragraphs),
                "content": "\n".join(paragraphs),
                "message": f"Document opened successfully"
            }
        except Exception as e:
            logger.error(f"Error opening document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def add_text(
        self,
        file_path: str,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        color: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Add text to a Word document
        
        Args:
            file_path: Path to the document
            text: Text to add
            bold: Make text bold
            italic: Make text italic
            underline: Underline text
            font_name: Font name (e.g., "Arial", "Times New Roman")
            font_size: Font size in points
            color: Text color (hex format, e.g., "#FF0000" for red)
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            para = doc.add_paragraph()
            run = para.add_run(text)
            
            # Apply formatting
            run.bold = bold
            run.italic = italic
            run.underline = underline
            
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            
            if font_size:
                run.font.size = Pt(font_size)
            
            if color:
                # Convert hex color to RGB
                color = color.lstrip('#')
                rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
                run.font.color.rgb = rgb
            
            doc.save(file_path)
            
            logger.info(f"Added text to document: {file_path}")
            return {
                "success": True,
                "message": "Text added successfully"
            }
        except Exception as e:
            logger.error(f"Error adding text: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def format_paragraph(
        self,
        file_path: str,
        paragraph_index: int,
        alignment: Optional[str] = None,
        line_spacing: Optional[float] = None,
        space_before: Optional[float] = None,
        space_after: Optional[float] = None,
        left_indent: Optional[float] = None,
        right_indent: Optional[float] = None
    ) -> Dict[str, Any]:
        """
        Format a paragraph in a Word document
        
        Args:
            file_path: Path to the document
            paragraph_index: Index of the paragraph to format (0-based)
            alignment: Paragraph alignment ("left", "center", "right", "justify")
            line_spacing: Line spacing (e.g., 1.5, 2.0)
            space_before: Space before paragraph in points
            space_after: Space after paragraph in points
            left_indent: Left indent in inches
            right_indent: Right indent in inches
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            
            if paragraph_index >= len(doc.paragraphs):
                return {
                    "success": False,
                    "error": f"Paragraph index {paragraph_index} out of range"
                }
            
            para = doc.paragraphs[paragraph_index]
            
            # Set alignment
            if alignment:
                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.lower() in align_map:
                    para.alignment = align_map[alignment.lower()]
            
            # Set line spacing
            if line_spacing is not None:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = line_spacing
            
            # Set spacing
            if space_before is not None:
                para.paragraph_format.space_before = Pt(space_before)
            if space_after is not None:
                para.paragraph_format.space_after = Pt(space_after)
            
            # Set indentation
            if left_indent is not None:
                para.paragraph_format.left_indent = Inches(left_indent)
            if right_indent is not None:
                para.paragraph_format.right_indent = Inches(right_indent)
            
            doc.save(file_path)
            
            logger.info(f"Formatted paragraph {paragraph_index} in document: {file_path}")
            return {
                "success": True,
                "message": "Paragraph formatted successfully"
            }
        except Exception as e:
            logger.error(f"Error formatting paragraph: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def add_heading(
        self,
        file_path: str,
        text: str,
        level: int = 1
    ) -> Dict[str, Any]:
        """
        Add a heading to a Word document
        
        Args:
            file_path: Path to the document
            text: Heading text
            level: Heading level (1-9)
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            doc.add_heading(text, level=min(level, 9))
            doc.save(file_path)
            
            logger.info(f"Added heading to document: {file_path}")
            return {
                "success": True,
                "message": "Heading added successfully"
            }
        except Exception as e:
            logger.error(f"Error adding heading: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def add_list(
        self,
        file_path: str,
        items: List[str],
        numbered: bool = False
    ) -> Dict[str, Any]:
        """
        Add a list (bulleted or numbered) to a Word document
        
        Args:
            file_path: Path to the document
            items: List of items to add
            numbered: True for numbered list, False for bulleted list
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            
            for item in items:
                para = doc.add_paragraph(item, style='List Number' if numbered else 'List Bullet')
            
            doc.save(file_path)
            
            logger.info(f"Added list to document: {file_path}")
            return {
                "success": True,
                "message": f"List with {len(items)} items added successfully"
            }
        except Exception as e:
            logger.error(f"Error adding list: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def add_table(
        self,
        file_path: str,
        rows: int,
        cols: int,
        data: Optional[List[List[str]]] = None,
        header_row: bool = False
    ) -> Dict[str, Any]:
        """
        Add a table to a Word document
        
        Args:
            file_path: Path to the document
            rows: Number of rows
            cols: Number of columns
            data: Optional 2D list of data to populate the table
            header_row: Whether the first row should be formatted as a header
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            table = doc.add_table(rows=rows, cols=cols)
            
            # Populate table with data if provided
            if data:
                for i, row_data in enumerate(data[:rows]):
                    for j, cell_data in enumerate(row_data[:cols]):
                        table.rows[i].cells[j].text = str(cell_data)
            
            # Format header row if requested
            if header_row and rows > 0:
                header_cells = table.rows[0].cells
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].bold = True
            
            doc.save(file_path)
            
            logger.info(f"Added table to document: {file_path}")
            return {
                "success": True,
                "message": f"Table with {rows}x{cols} dimensions added successfully"
            }
        except Exception as e:
            logger.error(f"Error adding table: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def find_replace(
        self,
        file_path: str,
        find_text: str,
        replace_text: str,
        replace_all: bool = True
    ) -> Dict[str, Any]:
        """
        Find and replace text in a Word document
        
        Args:
            file_path: Path to the document
            find_text: Text to find
            replace_text: Text to replace with
            replace_all: Whether to replace all occurrences
        
        Returns:
            Dictionary with success status and replacement count
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            replacement_count = 0
            
            for para in doc.paragraphs:
                if find_text in para.text:
                    para.text = para.text.replace(find_text, replace_text)
                    replacement_count += para.text.count(replace_text) if replace_all else 1
                    if not replace_all:
                        break
            
            # Also search in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find_text in cell.text:
                            cell.text = cell.text.replace(find_text, replace_text)
                            replacement_count += cell.text.count(replace_text) if replace_all else 1
                            if not replace_all:
                                break
            
            doc.save(file_path)
            
            logger.info(f"Find and replace completed in document: {file_path}")
            return {
                "success": True,
                "message": f"Replaced {replacement_count} occurrence(s)",
                "replacement_count": replacement_count
            }
        except Exception as e:
            logger.error(f"Error in find and replace: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def set_page_setup(
        self,
        file_path: str,
        margins: Optional[Dict[str, float]] = None,
        orientation: Optional[str] = None,
        page_size: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Set page setup options for a Word document
        
        Args:
            file_path: Path to the document
            margins: Dictionary with margin values in inches (top, bottom, left, right)
            orientation: Page orientation ("portrait" or "landscape")
            page_size: Page size (e.g., "Letter", "A4", "Legal")
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            from docx.enum.section import WD_ORIENT, WD_SECTION
            from docx.shared import Mm
            
            doc = Document(file_path)
            section = doc.sections[0]
            
            # Set margins
            if margins:
                if 'top' in margins:
                    section.top_margin = Inches(margins['top'])
                if 'bottom' in margins:
                    section.bottom_margin = Inches(margins['bottom'])
                if 'left' in margins:
                    section.left_margin = Inches(margins['left'])
                if 'right' in margins:
                    section.right_margin = Inches(margins['right'])
            
            # Set orientation
            if orientation:
                if orientation.lower() == "landscape":
                    section.orientation = WD_ORIENT.LANDSCAPE
                elif orientation.lower() == "portrait":
                    section.orientation = WD_ORIENT.PORTRAIT
            
            # Set page size
            if page_size:
                page_sizes = {
                    "Letter": (Inches(8.5), Inches(11)),
                    "A4": (Mm(210), Mm(297)),
                    "Legal": (Inches(8.5), Inches(14)),
                    "A3": (Mm(297), Mm(420)),
                    "A5": (Mm(148), Mm(210))
                }
                if page_size in page_sizes:
                    section.page_width, section.page_height = page_sizes[page_size]
            
            doc.save(file_path)
            
            logger.info(f"Page setup updated for document: {file_path}")
            return {
                "success": True,
                "message": "Page setup updated successfully"
            }
        except Exception as e:
            logger.error(f"Error setting page setup: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def save_document(self, file_path: str, new_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Save a Word document (or save as a new file)
        
        Args:
            file_path: Current path to the document
            new_path: Optional new path to save as
        
        Returns:
            Dictionary with success status
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            save_path = new_path or file_path
            
            # Ensure directory exists
            save_dir = os.path.dirname(save_path)
            if save_dir:
                os.makedirs(save_dir, exist_ok=True)
            
            doc.save(save_path)
            
            logger.info(f"Document saved: {save_path}")
            return {
                "success": True,
                "file_path": save_path,
                "message": f"Document saved successfully to {save_path}"
            }
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    async def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get information about a Word document
        
        Args:
            file_path: Path to the document
        
        Returns:
            Dictionary with document information
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed"
            }
        
        try:
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"Document not found: {file_path}"
                }
            
            doc = Document(file_path)
            
            # Get file stats
            file_stats = os.stat(file_path)
            
            info = {
                "success": True,
                "file_path": file_path,
                "file_size": file_stats.st_size,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "section_count": len(doc.sections),
                "created": file_stats.st_ctime,
                "modified": file_stats.st_mtime
            }
            
            # Get first few paragraphs as preview
            preview_paragraphs = [para.text for para in doc.paragraphs[:5]]
            info["preview"] = "\n".join(preview_paragraphs)
            
            return info
        except Exception as e:
            logger.error(f"Error getting document info: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    async def save_html_content(
        self,
        file_path: str,
        html_content: str
    ) -> Dict[str, Any]:
        """
        Save HTML content to a Word document, preserving formatting
        
        Args:
            file_path: Path where the document should be saved
            html_content: HTML content from contenteditable div
        
        Returns:
            Dictionary with success status and file path
        """
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx library not installed. Install with: pip install python-docx"
            }
        
        try:
            from html.parser import HTMLParser
            import re
            
            # Log the HTML content for debugging (first 500 chars)
            logger.info(f"Saving HTML content to {file_path}, HTML length: {len(html_content)}")
            logger.debug(f"HTML preview (first 2000 chars): {html_content[:2000]}")
            # Check if HTML contains tables
            if '<table' in html_content.lower():
                logger.info("HTML contains table(s), will parse with table support")
                # Count tables
                table_count = html_content.lower().count('<table')
                logger.info(f"Found {table_count} table(s) in HTML")
                # Extract table HTML for debugging
                import re
                table_matches = re.findall(r'<table[^>]*>.*?</table>', html_content, re.DOTALL | re.IGNORECASE)
                for i, table_html in enumerate(table_matches[:3]):  # Log first 3 tables
                    logger.debug(f"Table {i+1} HTML (first 500 chars): {table_html[:500]}")
            else:
                logger.warning("HTML does NOT contain any <table> tags!")
            
            # Normalize the file path
            original_path = file_path
            file_path = os.path.normpath(file_path)
            logger.info(f"Normalized path: {file_path} (original: {original_path})")
            
            # Ensure directory exists
            file_dir = os.path.dirname(file_path)
            if not file_dir:
                # If no directory specified, use default
                file_dir = self.default_documents_dir
                file_path = os.path.join(file_dir, os.path.basename(file_path))
                logger.warning(f"No directory in path, using default: {file_path}")
            
            logger.info(f"Creating directory if needed: {file_dir}")
            try:
                os.makedirs(file_dir, exist_ok=True)
                logger.info(f"Directory ready: {file_dir}")
            except Exception as dir_error:
                logger.error(f"Error creating directory {file_dir}: {dir_error}")
                return {
                    "success": False,
                    "error": f"Cannot create directory '{file_dir}': {str(dir_error)}"
                }
            
            logger.info(f"Final file path: {file_path}")
            
            # Create new document
            doc = Document()
            
            # Parse HTML and convert to Word document
            class HTMLToWordParser(HTMLParser):
                def __init__(self, doc):
                    super().__init__()
                    self.doc = doc
                    self.current_para = None
                    self.current_run = None
                    self.current_table = None
                    self.current_row = None
                    self.current_cell = None
                    self.table_cells_info = []
                    self.current_row_index = -1
                    self.current_cell_index = -1
                    self.in_table = False
                    
                def handle_starttag(self, tag, attrs):
                    attrs_dict = dict(attrs)
                    
                    if tag in ['p', 'div']:
                        # Start new paragraph
                        # If we're inside a table cell, use the cell's paragraph instead of creating a new one
                        if self.current_cell:
                            # Inside a cell, just ensure we have a paragraph
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                            # Reset run so new content creates a new run
                            self.current_run = None
                        else:
                            # Create a new paragraph in the document
                            self.current_para = self.doc.add_paragraph()
                            self.current_run = None
                    
                    elif tag == 'br':
                        # Line break
                        if self.current_cell:
                            # Inside a cell
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                            if self.current_run:
                                self.current_run.add_break()
                            else:
                                self.current_para.add_run().add_break()
                        elif self.current_para:
                            # In a paragraph, add break
                            if self.current_run:
                                self.current_run.add_break()
                            else:
                                self.current_para.add_run().add_break()
                        else:
                            # No paragraph context, create one with a break
                            self.current_para = self.doc.add_paragraph()
                            self.current_para.add_run().add_break()
                            self.current_run = None
                    
                    elif tag == 'strong' or tag == 'b':
                        # Bold
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        self.current_run.bold = True
                    
                    elif tag == 'em' or tag == 'i':
                        # Italic
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        self.current_run.italic = True
                    
                    elif tag == 'u':
                        # Underline
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        self.current_run.underline = True
                    
                    elif tag == 's' or tag == 'strike':
                        # Strikethrough
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        self.current_run.font.strike = True
                    
                    elif tag == 'sub':
                        # Subscript
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        self.current_run.font.subscript = True
                    
                    elif tag == 'sup':
                        # Superscript
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        self.current_run.font.superscript = True
                    
                    elif tag == 'a':
                        # Hyperlink
                        href = attrs_dict.get('href', '')
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                        elif self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        if self.current_run is None:
                            self.current_run = self.current_para.add_run()
                        # Store href for later use (python-docx hyperlinks need special handling)
                        self.current_run._hyperlink_url = href
                    
                    elif tag == 'ul' or tag == 'ol':
                        # Lists - handled in handle_data
                        if self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        self.current_run = None
                    
                    elif tag == 'li':
                        # List item
                        if self.current_para is None:
                            self.current_para = self.doc.add_paragraph()
                        self.current_run = None
                    
                    elif tag == 'table':
                        # Table - create new table (we'll add rows dynamically)
                        logger.info("=== PARSING TABLE ===")
                        logger.info(f"Table attributes: {attrs_dict}")
                        self.in_table = True
                        self.current_table = self.doc.add_table(rows=0, cols=0)
                        # Default style - will be overridden if HTML has style attribute
                        self.current_table.style = 'Light Grid Accent 1'
                        
                        # Parse table attributes and styles
                        table_style = attrs_dict.get('style', '')
                        table_class = attrs_dict.get('class', '')
                        logger.info(f"Table style: {table_style}, class: {table_class}")
                        
                        # Store table formatting info for later application
                        self.table_format_info = {
                            'style': table_style,
                            'class': table_class,
                            'border': attrs_dict.get('border', ''),
                            'cellpadding': attrs_dict.get('cellpadding', ''),
                            'cellspacing': attrs_dict.get('cellspacing', ''),
                            'width': attrs_dict.get('width', '')
                        }
                        
                        # Apply table width if specified
                        if 'width' in attrs_dict:
                            try:
                                width_str = attrs_dict['width']
                                if width_str.endswith('%'):
                                    # Percentage width - convert to inches (assuming 6.5" page width)
                                    width_pct = float(width_str.rstrip('%'))
                                    from docx.shared import Inches
                                    self.current_table.autofit = False
                                    # Set approximate width based on percentage
                                    self.current_table._element.tblPr.append(
                                        self.current_table._element._new_tblPr_child('w:tblW', {
                                            'w:w': str(int(width_pct * 50)),  # 50 = 0.5pt per 1%
                                            'w:type': 'pct'
                                        })
                                    )
                                elif width_str.endswith('px'):
                                    # Pixel width - approximate conversion
                                    width_px = float(width_str.rstrip('px'))
                                    from docx.shared import Inches
                                    width_inches = width_px / 96  # 96 DPI
                                    self.current_table.columns[0].width = Inches(width_inches) if len(self.current_table.columns) > 0 else None
                            except:
                                pass
                        
                        self.current_row = None
                        self.current_cell = None
                        self.current_para = None
                        self.current_run = None
                        self.table_cells_info = []  # Store cell info for post-processing
                        self.current_row_index = -1
                        self.current_cell_index = -1
                    
                    elif tag == 'tr':
                        # Table row
                        if self.current_table:
                            logger.debug(f"Adding table row {self.current_row_index + 1}")
                            self.current_row = self.current_table.add_row()
                            self.current_row_index += 1
                            self.current_cell_index = -1
                            self.current_cell = None
                            self.current_para = None
                            self.current_run = None
                        else:
                            logger.warning("Encountered <tr> tag but no current_table exists!")
                    
                    elif tag == 'td' or tag == 'th':
                        # Table cell
                        if self.current_row:
                            # Get colspan and rowspan from attributes
                            try:
                                colspan = int(attrs_dict.get('colspan', 1))
                            except:
                                colspan = 1
                            try:
                                rowspan = int(attrs_dict.get('rowspan', 1))
                            except:
                                rowspan = 1
                            
                            logger.debug(f"Adding cell at row {self.current_row_index}, col {self.current_cell_index + 1}, colspan={colspan}, rowspan={rowspan}")
                            
                            # Add cell to row
                            self.current_cell = self.current_row.add_cell()
                            self.current_cell_index += 1
                            
                            # Parse and apply cell styles
                            cell_style = attrs_dict.get('style', '')
                            
                            # Get paragraph from cell first (needed for alignment and content)
                            self.current_para = self.current_cell.paragraphs[0]
                            self.current_run = None
                            
                            # Apply cell background color
                            bg_color_match = re.search(r'background-color:\s*(#[0-9a-fA-F]{6}|rgb\([^)]+\)|\w+)', cell_style)
                            if bg_color_match:
                                try:
                                    bg_color_str = bg_color_match.group(1)
                                    if bg_color_str.startswith('#'):
                                        color_hex = bg_color_str.lstrip('#')
                                        # Set cell shading
                                        shading_elm = parse_xml(
                                            r'<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls('w'), color_hex)
                                        )
                                        tc_pr = self.current_cell._element.tcPr
                                        if tc_pr is None:
                                            tc_pr = self.current_cell._element._add_tcPr()
                                        existing_shd = tc_pr.find(qn('w:shd'))
                                        if existing_shd is not None:
                                            tc_pr.remove(existing_shd)
                                        tc_pr.append(shading_elm)
                                except Exception as e:
                                    logger.debug(f"Could not apply cell background color: {e}")
                            
                            # Apply cell text alignment
                            align_match = re.search(r'text-align:\s*(\w+)', cell_style)
                            if align_match:
                                align = align_match.group(1).lower()
                                from docx.enum.text import WD_ALIGN_PARAGRAPH
                                align_map = {
                                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                                }
                                if align in align_map:
                                    self.current_para.alignment = align_map[align]
                            
                            # Apply vertical alignment
                            vertical_align_match = re.search(r'vertical-align:\s*(\w+)', cell_style)
                            if vertical_align_match:
                                try:
                                    v_align = vertical_align_match.group(1).lower()
                                    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
                                    v_align_map = {
                                        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
                                        "middle": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                                        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                                        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                                    }
                                    if v_align in v_align_map:
                                        self.current_cell.vertical_alignment = v_align_map[v_align]
                                except:
                                    pass
                            
                            # Store cell info for later merging
                            cell_info = {
                                'cell': self.current_cell,
                                'row': self.current_row_index,
                                'col': self.current_cell_index,
                                'colspan': colspan,
                                'rowspan': rowspan,
                                'style': cell_style
                            }
                            self.table_cells_info.append(cell_info)
                        else:
                            logger.warning(f"Encountered <{tag}> tag but no current_row exists!")
                    
                    elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        # Heading
                        level = int(tag[1])
                        self.current_para = self.doc.add_heading(level=level)
                        self.current_run = None
                    
                    elif tag == 'span' or tag == 'font':
                        # Handle inline styles
                        # If we're in a table cell, use the cell's paragraph
                        if self.current_cell:
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                            if self.current_run is None:
                                self.current_run = self.current_para.add_run()
                        else:
                            if self.current_para is None:
                                self.current_para = self.doc.add_paragraph()
                            if self.current_run is None:
                                self.current_run = self.current_para.add_run()
                        
                        # Parse style attribute
                        style = attrs_dict.get('style', '')
                        if style:
                            # Parse font-size
                            font_size_match = re.search(r'font-size:\s*(\d+)pt', style)
                            if font_size_match:
                                self.current_run.font.size = Pt(int(font_size_match.group(1)))
                            
                            # Parse font-family
                            font_family_match = re.search(r'font-family:\s*([^;]+)', style)
                            if font_family_match:
                                font_name = font_family_match.group(1).strip().strip('"').strip("'").split(',')[0].strip()
                                self.current_run.font.name = font_name
                                self.current_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                            
                            # Parse color
                            color_match = re.search(r'color:\s*(#[0-9a-fA-F]{6}|rgb\([^)]+\))', style)
                            if color_match:
                                color_str = color_match.group(1)
                                if color_str.startswith('#'):
                                    # Hex color
                                    color = color_str.lstrip('#')
                                    rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
                                    self.current_run.font.color.rgb = rgb
                                elif color_str.startswith('rgb'):
                                    # RGB color
                                    rgb_match = re.search(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
                                    if rgb_match:
                                        rgb = RGBColor(int(rgb_match.group(1)), int(rgb_match.group(2)), int(rgb_match.group(3)))
                                        self.current_run.font.color.rgb = rgb
                            
                            # Parse background-color (highlight) - python-docx uses WD_COLOR_INDEX
                            bg_color_match = re.search(r'background-color:\s*(#[0-9a-fA-F]{6}|rgb\([^)]+\)|\w+)', style)
                            if bg_color_match:
                                bg_color_str = bg_color_match.group(1)
                                try:
                                    # Map common colors to WD_COLOR_INDEX
                                    color_map = {
                                        'yellow': WD_COLOR_INDEX.YELLOW,
                                        'green': WD_COLOR_INDEX.BRIGHT_GREEN,
                                        'blue': WD_COLOR_INDEX.BLUE,
                                        'red': WD_COLOR_INDEX.RED,
                                        'cyan': WD_COLOR_INDEX.TURQUOISE,
                                        'magenta': WD_COLOR_INDEX.PINK,
                                        'lightgray': WD_COLOR_INDEX.GRAY_25,
                                        'gray': WD_COLOR_INDEX.GRAY_50,
                                    }
                                    if bg_color_str.lower() in color_map:
                                        self.current_run.font.highlight_color = color_map[bg_color_str.lower()]
                                    elif bg_color_str.lower() == 'yellow' or '#ffff00' in bg_color_str.lower():
                                        self.current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                except:
                                    pass
                            
                            # Parse text-align
                            align_match = re.search(r'text-align:\s*(\w+)', style)
                            if align_match and self.current_para:
                                align = align_match.group(1).lower()
                                align_map = {
                                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                                    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                                }
                                if align in align_map:
                                    self.current_para.alignment = align_map[align]
                    
                    # Handle font tag attributes
                    if tag == 'font':
                        if 'size' in attrs_dict:
                            try:
                                size = int(attrs_dict['size'])
                                if self.current_para is None:
                                    self.current_para = self.doc.add_paragraph()
                                if self.current_run is None:
                                    self.current_run = self.current_para.add_run()
                                # Font size in points (approximate conversion)
                                self.current_run.font.size = Pt(size * 2 + 8)
                            except:
                                pass
                        if 'face' in attrs_dict:
                            if self.current_para is None:
                                self.current_para = self.doc.add_paragraph()
                            if self.current_run is None:
                                self.current_run = self.current_para.add_run()
                            self.current_run.font.name = attrs_dict['face']
                            self.current_run._element.rPr.rFonts.set(qn('w:eastAsia'), attrs_dict['face'])
                        if 'color' in attrs_dict:
                            if self.current_para is None:
                                self.current_para = self.doc.add_paragraph()
                            if self.current_run is None:
                                self.current_run = self.current_para.add_run()
                            color = attrs_dict['color'].lstrip('#')
                            if len(color) == 6:
                                rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
                                self.current_run.font.color.rgb = rgb
                
                def handle_endtag(self, tag):
                    if tag in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'ul', 'ol']:
                        # End paragraph
                        self.current_para = None
                        self.current_run = None
                    elif tag in ['strong', 'b', 'em', 'i', 'u', 's', 'strike', 'sub', 'sup', 'span', 'font', 'a']:
                        # End run - create new run for next text
                        self.current_run = None
                    elif tag == 'td' or tag == 'th':
                        # End table cell
                        self.current_cell = None
                        self.current_para = None
                        self.current_run = None
                    elif tag == 'tr':
                        # End table row
                        self.current_row = None
                        self.current_cell = None
                        self.current_para = None
                        self.current_run = None
                    elif tag == 'table':
                        # End table - apply table formatting and process merges
                        if self.current_table:
                            # Apply table borders and formatting if specified
                            if hasattr(self, 'table_format_info'):
                                format_info = self.table_format_info
                                # Check if border is specified in style or attribute
                                border_specified = False
                                if format_info.get('border') or 'border' in format_info.get('style', ''):
                                    border_specified = True
                                
                                # Apply borders to all cells if border is specified
                                if border_specified:
                                    try:
                                        from docx.oxml import parse_xml
                                        # Set table borders
                                        tbl_borders = parse_xml(
                                            r'<w:tblBorders {}>'
                                            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                                            r'</w:tblBorders>'.format(nsdecls('w'))
                                        )
                                        tbl_pr = self.current_table._element.tblPr
                                        existing_borders = tbl_pr.find(qn('w:tblBorders'))
                                        if existing_borders is not None:
                                            tbl_pr.remove(existing_borders)
                                        tbl_pr.append(tbl_borders)
                                    except Exception as e:
                                        logger.debug(f"Could not apply table borders: {e}")
                                
                                # Clean up format info
                                if hasattr(self, 'table_format_info'):
                                    delattr(self, 'table_format_info')
                        
                        # Process colspan and rowspan merges
                        if self.current_table and hasattr(self, 'table_cells_info'):
                            try:
                                # Process merges: first handle colspan, then rowspan
                                rows = list(self.current_table.rows)
                                
                                # Build a map of cells by position
                                cell_map = {}
                                for info in self.table_cells_info:
                                    row_idx = info['row']
                                    col_idx = info['col']
                                    if row_idx < len(rows):
                                        row = rows[row_idx]
                                        if col_idx < len(row.cells):
                                            cell_map[(row_idx, col_idx)] = {
                                                'cell': row.cells[col_idx],
                                                'colspan': info['colspan'],
                                                'rowspan': info['rowspan']
                                            }
                                
                                # Process merges - sort by row then column to process in order
                                sorted_cells = sorted(cell_map.items(), key=lambda x: (x[0][0], x[0][1]))
                                
                                for (row_idx, col_idx), cell_data in sorted_cells:
                                    cell = cell_data['cell']
                                    colspan = cell_data['colspan']
                                    rowspan = cell_data['rowspan']
                                    
                                    # Skip if cell has already been merged (removed from row)
                                    try:
                                        # Verify cell still exists in the row
                                        if row_idx >= len(rows) or col_idx >= len(rows[row_idx].cells):
                                            continue
                                        if rows[row_idx].cells[col_idx] != cell:
                                            continue
                                    except:
                                        continue
                                    
                                    # Handle colspan (merge cells horizontally)
                                    if colspan > 1:
                                        try:
                                            for c in range(1, colspan):
                                                merge_col = col_idx + c
                                                if merge_col < len(rows[row_idx].cells):
                                                    next_cell = rows[row_idx].cells[merge_col]
                                                    try:
                                                        cell.merge(next_cell)
                                                        logger.debug(f"Merged colspan: ({row_idx}, {col_idx}) with ({row_idx}, {merge_col})")
                                                    except Exception as merge_err:
                                                        # Cell might already be merged
                                                        logger.debug(f"Could not merge colspan: {merge_err}")
                                        except Exception as e:
                                            logger.warning(f"Error merging colspan for cell at ({row_idx}, {col_idx}): {e}")
                                    
                                    # Handle rowspan (merge cells vertically)
                                    if rowspan > 1:
                                        try:
                                            for r in range(1, rowspan):
                                                merge_row_idx = row_idx + r
                                                if merge_row_idx < len(rows):
                                                    merge_row = rows[merge_row_idx]
                                                    # Find cell in same column (accounting for previous horizontal merges)
                                                    merge_col = col_idx
                                                    if merge_col < len(merge_row.cells):
                                                        cell_to_merge = merge_row.cells[merge_col]
                                                        try:
                                                            cell.merge(cell_to_merge)
                                                            logger.debug(f"Merged rowspan: ({row_idx}, {col_idx}) with ({merge_row_idx}, {merge_col})")
                                                        except Exception as merge_err:
                                                            # Cell might already be merged
                                                            logger.debug(f"Could not merge rowspan: {merge_err}")
                                        except Exception as e:
                                            logger.warning(f"Error merging rowspan for cell at ({row_idx}, {col_idx}): {e}")
                                
                                # Clean up
                                if hasattr(self, 'table_cells_info'):
                                    del self.table_cells_info
                            except Exception as e:
                                logger.error(f"Error processing table merges: {e}")
                                import traceback
                                logger.error(traceback.format_exc())
                        
                        self.in_table = False
                        self.current_table = None
                        self.current_row = None
                        self.current_cell = None
                        self.current_para = None
                        self.current_run = None
                        self.current_row_index = -1
                        self.current_cell_index = -1
                        if hasattr(self, 'table_cells_info'):
                            self.table_cells_info = []
                
                def handle_data(self, data):
                    # Add text to current run
                    try:
                        # Decode HTML entities
                        import html
                        data = html.unescape(data)
                        
                        # Process all data - ensure we have a paragraph and run before adding text
                        # If we're in a table cell, always add content (even whitespace)
                        if self.current_cell:
                            # Ensure we have a paragraph in the cell
                            if self.current_para is None:
                                self.current_para = self.current_cell.paragraphs[0]
                            # Ensure we have a run
                            if self.current_run is None:
                                self.current_run = self.current_para.add_run()
                            # Add data - replace &nbsp; with regular space for better compatibility
                            text_to_add = data.replace('\xa0', ' ')  # Replace non-breaking space with regular space
                            # Always add text if we have a run (even if it's whitespace)
                            self.current_run.add_text(text_to_add)
                        # If we have a paragraph context, add the data
                        elif self.current_para:
                            # Ensure we have a run
                            if self.current_run is None:
                                self.current_run = self.current_para.add_run()
                            text_to_add = data.replace('\xa0', ' ')
                            # Add text (even whitespace to preserve formatting)
                            self.current_run.add_text(text_to_add)
                        # If we have non-whitespace data and no context, create a paragraph
                        elif data.strip():
                            # Create paragraph and run for new content
                            self.current_para = self.doc.add_paragraph()
                            self.current_run = self.current_para.add_run()
                            text_to_add = data.replace('\xa0', ' ')
                            self.current_run.add_text(text_to_add)
                        # For pure whitespace with no context, we can skip it
                        # But if we're in the middle of parsing, we might want to preserve it
                        # For now, skip only if we truly have no context
                        else:
                            # Skip pure whitespace with no context
                            return
                    except Exception as e:
                        logger.warning(f"Error handling data in HTML parser: {e}")
                        import traceback
                        logger.warning(traceback.format_exc())
                        # Try to add text anyway
                        if self.current_cell:
                            try:
                                if self.current_para is None:
                                    self.current_para = self.current_cell.paragraphs[0]
                                if self.current_run is None:
                                    self.current_run = self.current_para.add_run()
                                self.current_run.add_text(str(data))
                            except Exception as e2:
                                logger.warning(f"Failed to add text to cell: {e2}")
                        else:
                            try:
                                if self.current_para is None:
                                    self.current_para = self.doc.add_paragraph()
                                if self.current_run is None:
                                    self.current_run = self.current_para.add_run()
                                self.current_run.add_text(str(data))
                            except Exception as e2:
                                logger.warning(f"Failed to add text to paragraph: {e2}")
            
            # Handle empty or minimal content
            if not html_content or not html_content.strip():
                # If content is empty, add at least one paragraph
                doc.add_paragraph()
            else:
                # Parse HTML content
                parser = HTMLToWordParser(doc)
                try:
                    # Clean up HTML content - remove script and style tags
                    import re
                    cleaned_html = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                    cleaned_html = re.sub(r'<style[^>]*>.*?</style>', '', cleaned_html, flags=re.DOTALL | re.IGNORECASE)
                    
                    # Check for tables in HTML before parsing
                    table_count_in_html = len(re.findall(r'<table[^>]*>', cleaned_html, re.IGNORECASE))
                    logger.info(f"Found {table_count_in_html} table(s) in HTML content")
                    
                    # Log cleaned HTML for debugging
                    logger.info(f"Cleaned HTML length: {len(cleaned_html)}")
                    if len(cleaned_html) < 1000:
                        logger.debug(f"Full cleaned HTML: {cleaned_html}")
                    else:
                        logger.debug(f"Cleaned HTML preview (first 1000 chars): {cleaned_html[:1000]}")
                        # Also log table sections specifically
                        table_matches = re.findall(r'<table[^>]*>.*?</table>', cleaned_html, re.DOTALL | re.IGNORECASE)
                        for idx, table_html in enumerate(table_matches):
                            logger.debug(f"Table {idx + 1} HTML (first 500 chars): {table_html[:500]}")
                    
                    logger.info("Starting HTML parsing...")
                    parser.feed(cleaned_html)
                    parser.close()  # Ensure parser is closed
                    logger.info("HTML parsing completed")
                    
                    # Log parsing results
                    logger.info(f"Parsing complete. Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
                    
                    # Verify tables were created correctly
                    if len(doc.tables) > 0:
                        logger.info(f"=== TABLE VERIFICATION ===")
                        for idx, table in enumerate(doc.tables):
                            logger.info(f"Table {idx + 1}: {len(table.rows)} rows, {len(table.columns) if table.rows else 0} columns")
                            for row_idx, row in enumerate(table.rows):
                                logger.info(f"  Row {row_idx + 1}: {len(row.cells)} cells")
                                for cell_idx, cell in enumerate(row.cells):
                                    cell_text = ' '.join([para.text for para in cell.paragraphs])
                                    logger.debug(f"    Cell ({row_idx}, {cell_idx}): '{cell_text[:50]}...' (length: {len(cell_text)})")
                    
                    # Check if we actually got content
                    total_text = sum(len(p.text) for p in doc.paragraphs)
                    # Also check text in table cells
                    table_text = 0
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    table_text += len(para.text)
                    total_content = total_text + table_text
                    logger.info(f"Total text content: {total_text} chars in paragraphs, {table_text} chars in tables, {total_content} total")
                    
                    if total_content == 0:
                        logger.error("WARNING: No content was parsed from HTML! Document will be empty.")
                        # IMPORTANT: If we have tables in HTML but no tables were created, don't use plain text fallback
                        # as it would destroy the table structure. Instead, log the error and let the user know.
                        if table_count_in_html > 0:
                            logger.error("CRITICAL: Tables were detected in HTML but parsing failed completely!")
                            logger.error("The document will be empty. This is a parser error that needs to be fixed.")
                            # Add an error message paragraph so the document isn't completely empty
                            doc.add_paragraph(f"ERROR: Table content could not be parsed. {table_count_in_html} table(s) were detected in the HTML but could not be converted to Word format.")
                        else:
                            # Only use plain text fallback if there were no tables
                            logger.warning("No tables detected, using plain text fallback")
                            import html
                            plain_text = html.unescape(cleaned_html)
                            # Remove HTML tags but preserve structure hints
                            plain_text = re.sub(r'<[^>]+>', ' ', plain_text)
                            # Clean up multiple spaces
                            plain_text = re.sub(r'\s+', ' ', plain_text).strip()
                            if plain_text:
                                logger.warning(f"Adding fallback plain text: {plain_text[:100]}...")
                                # Split by common separators and add as paragraphs
                                lines = plain_text.split('\n')
                                for line in lines:
                                    if line.strip():
                                        doc.add_paragraph(line.strip())
                                # If no lines, add the whole text
                                if len(doc.paragraphs) == 0:
                                    doc.add_paragraph(plain_text)
                    
                    if len(doc.tables) == 0 and table_count_in_html > 0:
                        logger.error(f"CRITICAL ERROR: HTML contained {table_count_in_html} table(s) but NO tables were created in the Word document!")
                        logger.error("This means table parsing failed. The document will contain only text, not table structures.")
                        logger.error("Please check the HTML structure and parser logic.")
                    elif len(doc.tables) > 0 and table_count_in_html > 0:
                        logger.info(f"SUCCESS: Created {len(doc.tables)} table(s) from {table_count_in_html} table(s) in HTML")
                except Exception as parse_error:
                    logger.error(f"HTML parsing error: {parse_error}")
                    import traceback
                    logger.error(traceback.format_exc())
                    
                    # Check if we have tables in the document despite the error
                    if len(doc.tables) > 0:
                        logger.info(f"Despite parsing error, {len(doc.tables)} table(s) were created. Continuing with saved tables...")
                    elif '<table' in html_content.lower():
                        logger.error("CRITICAL: HTML contained tables but parsing failed and no tables were created!")
                        logger.error("This will result in table content being lost. Check the error above.")
                        # Don't fall back to plain text - that would lose the table structure
                        # Instead, try to preserve what we can - at least save the document with what we have
                        logger.warning("Attempting to preserve table structure...")
                        # Add a note about the parsing error
                        doc.add_paragraph("Note: Table content could not be parsed correctly. Please check the logs.")
                    else:
                        # Only fall back to plain text if there were no tables expected
                        logger.warning("No tables in HTML, falling back to plain text extraction")
                        import html
                        plain_text = html.unescape(html_content)
                        # Remove HTML tags
                        plain_text = re.sub(r'<[^>]+>', '', plain_text)
                        if plain_text.strip():
                            doc.add_paragraph(plain_text.strip())
                        else:
                            doc.add_paragraph()
                    
                    # Continue with save even if parsing had errors (unless it's a critical error)
                    logger.info("Continuing with save despite parsing errors...")
            
            # Ensure at least one paragraph exists
            if len(doc.paragraphs) == 0:
                doc.add_paragraph()
            
            # Save document
            try:
                logger.info(f"Attempting to save document to: {file_path}")
                logger.info(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
                doc.save(file_path)
                logger.info(f"Successfully saved HTML content to Word document: {file_path}")
                
                # Verify file was created
                if os.path.exists(file_path):
                    file_size = os.path.getsize(file_path)
                    logger.info(f"File verified: {file_path} ({file_size} bytes)")
                else:
                    logger.error(f"File was not created at {file_path}")
                    return {
                        "success": False,
                        "error": f"File was not created at {file_path}"
                    }
                
                return {
                    "success": True,
                    "file_path": file_path,
                    "message": f"Document saved successfully at {file_path}"
                }
            except PermissionError as pe:
                error_msg = f"Permission denied: Cannot save to '{file_path}'. "
                error_msg += "You may not have write permissions for this location. "
                error_msg += "Try saving to your Documents folder instead."
                logger.error(f"Permission error saving to {file_path}: {pe}")
                return {
                    "success": False,
                    "error": error_msg
                }
            except OSError as ose:
                if ose.errno == 13:  # Permission denied
                    error_msg = f"Permission denied: Cannot save to '{file_path}'. "
                    error_msg += "You may not have write permissions for this location. "
                    error_msg += "Try saving to your Documents folder instead."
                    logger.error(f"Permission error saving to {file_path}: {ose}")
                    return {
                        "success": False,
                        "error": error_msg
                    }
                else:
                    logger.error(f"OS error saving HTML content: {ose}")
                    import traceback
                    logger.error(traceback.format_exc())
                    return {
                        "success": False,
                        "error": f"Error saving file: {str(ose)}"
                    }
        except Exception as e:
            logger.error(f"Error saving HTML content: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return {
                "success": False,
                "error": str(e)
            }

